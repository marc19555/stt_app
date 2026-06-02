import os
import sys
import json

sys.path.append(os.path.dirname(__file__))
import database as db

def export_all(session_id, folder, merged_path, pv_path, summary_path):
    output_folder = os.path.join(folder, 'outputs')
    os.makedirs(output_folder, exist_ok=True)

    
    pv_output_path = os.path.join(output_folder, 'pv.docx')
    _export_docx(pv_path, pv_output_path, 'Procès-Verbal')
    db.save_artifact(session_id, 'pv_docx', os.path.abspath(pv_output_path))

    summary_output_path = os.path.join(output_folder, 'resumer.docx')
    _export_docx(summary_path, summary_output_path, 'Resumer')
    db.save_artifact(session_id, 'resumer_docx', os.path.abspath(summary_output_path))

def _export_docx(md_path, docx_path, title):
    from docx import Document

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()
    doc.add_heading(title, level=0)

    for line in lines:
        line = line.rstrip()
        if line.startswith('## '):
            doc.add_heading(line[3:], level=1)
        elif line.startswith('# '):
            doc.add_heading(line[2:], level=0)
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('|'):
            doc.add_paragraph(line)
        elif line.strip() in ('', '---'):
            doc.add_paragraph('')
        else:
            doc.add_paragraph(line)

    doc.save(docx_path)
    print(f"{os.path.basename(docx_path)} Créé")
